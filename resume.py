from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Helper function to add a clickable hyperlink
def add_hyperlink(paragraph, url, text, color=RGBColor(0, 0, 255), underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Color element
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color.rgb.hex)
    rPr.append(c)

    # Underline element
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single' if underline else 'none')
    rPr.append(u)

    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

# Create the document
doc = Document()

# Define styles
def set_font(run, name='Arial', size=12, color=RGBColor(0, 0, 0), bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name)

# HEADER
header = doc.sections[0].header
header_para = header.paragraphs[0]
header_para.text = "Hayredin Mohammed"
header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = header_para.runs[0]
set_font(run, size=24, bold=True, color=RGBColor(0, 51, 102))

# Contact info below header
contact = doc.add_paragraph()
contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
contact.add_run("Addis Ababa, Ethiopia | ").bold = False
contact.add_run("📧 hayredin.950@gmail.com | ").bold = False
contact.add_run("📞 +251 940 522 137 | ").bold = False
contact.add_run("🔗 LinkedIn: linkedin.com/in/hayredin-mohammed-43a128311 | ").bold = False
contact.add_run("🐙 GitHub: github.com/HayreKhan750").bold = False

doc.add_paragraph()

# OBJECTIVE
obj_title = doc.add_paragraph()
obj_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
obj_run = obj_title.add_run("🎯 Objective")
set_font(obj_run, size=16, bold=True, color=RGBColor(0, 102, 204))
doc.add_paragraph(
    "Motivated Computer Science student with a solid foundation in Python, Machine Learning, and Data Analysis, "
    "seeking an internship position to apply and expand skills in real-world ML projects. "
    "Passionate about solving impactful problems through AI."
)

doc.add_paragraph()

# EDUCATION
edu_title = doc.add_paragraph()
edu_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
edu_run = edu_title.add_run("🎓 Education")
set_font(edu_run, size=16, bold=True, color=RGBColor(0, 102, 204))
doc.add_paragraph(
    "Bachelor of Science in Computer Science\n"
    "Addis Ababa University — Addis Ababa, Ethiopia\n"
    "Expected Graduation: June 2027\n\n"
    "Relevant Coursework: Machine Learning, Data Science, Python Programming, Statistics, Linear Algebra"
)

doc.add_paragraph()

# TECHNICAL SKILLS
skills_title = doc.add_paragraph()
skills_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
skills_run = skills_title.add_run("🛠️ Technical Skills")
set_font(skills_run, size=16, bold=True, color=RGBColor(0, 102, 204))
skills = [
    "Programming: Python, C++, SQL, Java (basic)",
    "Machine Learning & Data: Scikit-learn, Pandas, NumPy, Matplotlib",
    "Tools: Git, GitHub, Jupyter Notebook, VS Code",
    "Concepts: Data Preprocessing, Model Training & Evaluation, NLP, Sentiment Analysis"
]
for skill in skills:
    p = doc.add_paragraph(skill, style='List Bullet')

doc.add_paragraph()

# PROJECTS
proj_title = doc.add_paragraph()
proj_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
proj_run = proj_title.add_run("📊 Projects")
set_font(proj_run, size=16, bold=True, color=RGBColor(0, 102, 204))

# Project 1 - Sales Data Analysis
doc.add_paragraph("Sales Data Analysis")
p1_link = doc.add_paragraph("GitHub: ")
add_hyperlink(p1_link, "https://github.com/HayreKhan750/Sales-Data-Analysis", "https://github.com/HayreKhan750/Sales-Data-Analysis")
proj1_points = [
    "Performed exploratory data analysis to uncover insights and trends from sales datasets.",
    "Visualized key sales metrics using Matplotlib and Seaborn to support business decisions.",
    "Applied data cleaning and preprocessing to prepare datasets for analysis."
]
for point in proj1_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_paragraph()

# Project 2 - Twitter Sentiment Analyzer
doc.add_paragraph("Twitter Sentiment Analyzer")
p2_link = doc.add_paragraph("GitHub: ")
add_hyperlink(p2_link, "https://github.com/HayreKhan750/twitter-sentiment-analyzer", "https://github.com/HayreKhan750/twitter-sentiment-analyzer")
proj2_points = [
    "Built a sentiment analysis tool to classify tweets as positive, negative, or neutral using NLP techniques.",
    "Collected and cleaned Twitter data; performed tokenization, stopword removal, and feature extraction.",
    "Trained and evaluated models using scikit-learn classifiers (e.g., Naive Bayes, SVM).",
    "Visualized results with Matplotlib for clear insights into sentiment trends."
]
for point in proj2_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_paragraph()

# ACHIEVEMENTS & CERTIFICATES
achieve_title = doc.add_paragraph()
achieve_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
achieve_run = achieve_title.add_run("🏅 Achievements & Certifications")
set_font(achieve_run, size=16, bold=True, color=RGBColor(0, 102, 204))
achievements = [
    "Awarded Scholarship by Kibur College for academic excellence.",
    "Udacity Nanodegree Certificates: Android Fundamentals, Data Analysis, Programming Fundamentals."
]
for achievement in achievements:
    doc.add_paragraph(achievement, style='List Bullet')

doc.add_paragraph()

# LANGUAGES
lang_title = doc.add_paragraph()
lang_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
lang_run = lang_title.add_run("🌐 Languages")
set_font(lang_run, size=16, bold=True, color=RGBColor(0, 102, 204))
languages = [
    "English (Fluent)",
    "Amharic (Native)"
]
for lang in languages:
    doc.add_paragraph(lang, style='List Bullet')

doc.add_paragraph()

# REFERENCES
ref_title = doc.add_paragraph()
ref_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
ref_run = ref_title.add_run("🤝 References")
set_font(ref_run, size=16, bold=True, color=RGBColor(0, 102, 204))
doc.add_paragraph("Available upon request.")

# Save the document
file_path = "Hayredin_Mohammed_ML_Internship.docx"
doc.save(file_path)

print(f"Saved resume as: {file_path}")
